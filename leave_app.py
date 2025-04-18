
import streamlit as st
import pandas as pd
from datetime import date
import uuid
import os
import time
import requests
from docx import Document

# Load employee DB
df = pd.read_csv("leave_db.csv")

st.title("📝 Leave Application Form")

# --- Select Employee ---
selected_name = st.selectbox("Select Employee", df['name'])
emp = df[df['name'] == selected_name].iloc[0]

# --- Autofill Section ---
st.subheader("Employee Information")
st.text_input("Email Address", emp['email'], disabled=True)
st.text_input("PMCC Entity", emp['entity'], disabled=True)
st.text_input("Employee No.", emp['employee_no'], disabled=True)
st.text_input("Contact No.", emp['contact_no'], disabled=True)
st.date_input("Date", date.today(), disabled=True)

# --- Leave Info ---
st.subheader("Leave Information")
category = st.selectbox("Category", ["Annual", "Medical"])
start_date = st.date_input("From")
end_date = st.date_input("To")

# Calculate days
if start_date and end_date:
    total_days = (end_date - start_date).days + 1
else:
    total_days = 0
st.number_input("Total Number of Days", value=total_days, disabled=True)

day_type = st.radio("Day Type", ["Full Day", "Half Day - AM", "Half Day - PM"])
reason = st.text_area("Reason for Application")
attachment = st.file_uploader("Attachment (e.g. MC, cert)")

# --- Submit ---
if st.button("Submit"):
    leave_id = f"L{str(uuid.uuid4())[:8].upper()}"
    history_path = "leave_history.csv"

    # Handle file saving
    upload_dir = "uploads"
    os.makedirs(upload_dir, exist_ok=True)
    if attachment:
        timestamp = int(time.time())
        ext = os.path.splitext(attachment.name)[1]
        saved_file_name = f"{leave_id}_{timestamp}{ext}"
        file_path = os.path.join(upload_dir, saved_file_name)
        with open(file_path, "wb") as f:
            f.write(attachment.getbuffer())
    else:
        saved_file_name = ""

    # Create record
    record = {
        "leave_id": leave_id,
        "employee_id": emp['employee_id'],
        "leave_type": category,
        "from_date": start_date,
        "to_date": end_date,
        "total_days": total_days,
        "day_type": day_type,
        "reason": reason,
        "attachment": saved_file_name,
        "status": "Pending",
        "checker_name": "",
        "approver_name": ""
    }

    # Append to leave_history.csv
    if os.path.exists(history_path):
        history_df = pd.read_csv(history_path)
        history_df = history_df.append(record, ignore_index=True)
    else:
        history_df = pd.DataFrame([record])

    history_df.to_csv(history_path, index=False)

    # Trigger n8n webhook
    webhook_url = "https://your-n8n-instance/webhook/leave-approval"  # Replace with real URL
    try:
        requests.post(webhook_url, json=record)
    except Exception as e:
        st.warning(f"Could not notify approval system: {e}")

    # Generate DOCX leave form
    doc = Document()
    doc.add_heading("Leave Application Form", 0)
    doc.add_paragraph(f"Leave ID: {leave_id}")
    doc.add_paragraph(f"Name: {emp['name']}")
    doc.add_paragraph(f"Employee No.: {emp['employee_no']}")
    doc.add_paragraph(f"Contact No.: {emp['contact_no']}")
    doc.add_paragraph(f"Department: {emp['department']}")
    doc.add_paragraph(f"Entity: {emp['entity']}")
    doc.add_paragraph(f"Email: {emp['email']}")
    doc.add_paragraph(f"Leave Type: {category}")
    doc.add_paragraph(f"From: {start_date}")
    doc.add_paragraph(f"To: {end_date}")
    doc.add_paragraph(f"Total Days: {total_days}")
    doc.add_paragraph(f"Day Type: {day_type}")
    doc.add_paragraph(f"Reason: {reason}")
    doc.add_paragraph(f"Attachment: {saved_file_name}")
    doc.add_paragraph(f"Status: Pending")

    doc_path = os.path.join("uploads", f"{leave_id}_form.docx")
    doc.save(doc_path)

    st.success("Leave application submitted and document generated successfully!")
